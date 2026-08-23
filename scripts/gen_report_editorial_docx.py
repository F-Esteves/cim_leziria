import json
import subprocess
import shutil
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
CHARTS = BASE_DIR / "reports" / "charts_standalone"
OUT_DIR = BASE_DIR / "reports"
OUT_DIR.mkdir(parents=True, exist_ok=True)

with open(CHARTS / "narrativas.json", encoding="utf-8") as f:
    narrativas = json.load(f)

resumo_path = CHARTS / "resumo_indicadores.json"
resumo_indicadores = json.loads(resumo_path.read_text(encoding="utf-8")) if resumo_path.exists() else []

with open(CHARTS / "tabela_municipios.json", encoding="utf-8") as f:
    tabela_dados = json.load(f)

doc = Document()

# ── Margens reduzidas (mais área útil por página) ──
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ── Estilos globais ──
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(10)

HEADING_COLORS = {
    "Heading 1": (RGBColor(0x1F, 0x4E, 0x79), Pt(18), True),
    "Heading 2": (RGBColor(0x2E, 0x75, 0xB6), Pt(14), True),
    "Heading 3": (RGBColor(0x1F, 0x4E, 0x79), Pt(12), True),
}
for style_name, (color, size, bold) in HEADING_COLORS.items():
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st.font.color.rgb = color
    st.font.size = size
    st.font.bold = bold


# ── Helpers ──

def adicionar_quebra_pagina():
    doc.add_page_break()


def titulo_seccao(texto, com_quebra_pagina=True):
    if com_quebra_pagina:
        adicionar_quebra_pagina()
    h = doc.add_heading(texto, level=1)
    h.paragraph_format.space_after = Pt(8)


def sub_titulo(texto):
    h = doc.add_heading(texto, level=2)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(6)


def _shading_celula(cell, cor_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _bordas_celula(cell, bottom=True, cor="D9D9D9", tamanho=6):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    if bottom:
        bottom_el = OxmlElement("w:bottom")
        bottom_el.set(qn("w:val"), "single")
        bottom_el.set(qn("w:sz"), str(tamanho))
        bottom_el.set(qn("w:color"), cor)
        tcBorders.append(bottom_el)
    tcPr.append(tcBorders)


def bloco_simples(titulo, chave_imagem, chave_narrativa, opts=None):
    opts = opts or {}

    if titulo:
        p = doc.add_heading(titulo, level=3)
        p.paragraph_format.space_before = Pt(6 if opts.get("semEspacoAntes") else 16)
        p.paragraph_format.space_after = Pt(6)

    img_path = CHARTS / f"{chave_imagem}.png"
    width_px = opts.get("width", 400)
    width_cm = Cm(width_px / 96 * 2.54)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(img_path), width=width_cm)
    p_img.paragraph_format.space_after = Pt(6)

    texto = narrativas.get(chave_narrativa)
    if texto:
        p_txt = doc.add_paragraph(texto)
        p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_txt.paragraph_format.space_after = Pt(12)


def tabela_municipios():
    municipios = tabela_dados["municipios"]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0].cells
    headers = ["Município", "População", "Densidade (hab./km²)"]
    for i, texto in enumerate(headers):
        header[i].text = texto
        for p in header[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shading_celula(header[i], "1F4E79")

    for m in municipios:
        row = table.add_row().cells
        row[0].text = m["municipio"]
        row[1].text = f'{m["populacao"]:,}'.replace(",", ".")
        row[2].text = str(m["densidade"])
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


def tabela_resumo():
    """Tabela-síntese tipo 'semáforo': mostra, para cada indicador selecionado,
    a tendência (▲ aumento / ▼ diminuição / ● estável) e o cumprimento da meta
    (quando existe). Os dados vêm todos de resumo_indicadores.json, gerado
    automaticamente a partir da mesma classificação (narrativa_engine.
    avaliar_indicador) que produz o texto do relatório — não há aqui nenhuma
    avaliação feita à mão.

    A cor é aplicada só às colunas "Tendência" e "Meta" (não à linha inteira),
    de propósito: são dois sinais independentes — um indicador pode ter uma
    tendência positiva e ainda assim não cumprir a meta (ou vice-versa).
    Colorir a linha toda obrigaria a escolher um veredito único mesmo quando
    os dois sinais dizem coisas diferentes, escondendo informação real.
    """
    if not resumo_indicadores:
        return

    doc.add_page_break()

    VERDE = ("E2F0D9", "375623")
    VERMELHO = ("FCE4E4", "8B0000")
    CINZA = ("F2F2F2", "595959")
    SETAS = {"aumento": "▲", "diminuicao": "▼", "estabilidade": "●", "sem_comparacao": "—"}

    doc.add_heading("Síntese dos Indicadores", level=1)
    p = doc.add_paragraph(
        "Classificação automática da tendência e do cumprimento de metas para um conjunto de indicadores "
        "representativo de cada domínio, gerada a partir das mesmas regras (narrativa_engine.py) que produzem "
        "o texto do relatório — não há aqui avaliação qualitativa feita manualmente."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (Cm(2.6), Cm(6.0), Cm(2.6), Cm(2.2), Cm(2.9))
    headers = ["Domínio", "Indicador", "Valor Atual", "Tendência", "Meta"]

    header_cells = table.rows[0].cells
    for i, (texto, w) in enumerate(zip(headers, widths)):
        header_cells[i].width = w
        header_cells[i].text = texto
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shading_celula(header_cells[i], "1F4E79")

    for item in resumo_indicadores:
        row = table.add_row().cells
        for cel, w in zip(row, widths):
            cel.width = w

        row[0].text = item["cluster"]
        row[1].text = item["nome"]

        valor_fmt = f'{item["valor_atual"]:,.1f}'.replace(",", " ").rstrip("0").rstrip(".")
        row[2].text = f'{valor_fmt}{item["unidade"]}'
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Coluna Tendência: cor consoante a tendência em si for boa ou má
        # notícia (independente da meta). Estabilidade/sem-comparação/sem
        # direção definida ficam neutras (cinza) — não há "sinal" a dar.
        situacao = item["situacao"]
        direcao_boa = item.get("config", {}).get("direcao_boa")
        row[3].text = SETAS.get(situacao, "—")
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if situacao in ("aumento", "diminuicao") and direcao_boa:
            cor_tendencia = VERDE if situacao == direcao_boa else VERMELHO
        else:
            cor_tendencia = CINZA
        _shading_celula(row[3], cor_tendencia[0])
        for run in row[3].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(cor_tendencia[1])

        # Coluna Meta: cor consoante cumpre/não cumpre, independente da
        # tendência. Sem meta definida fica neutra (cinza).
        situacao_meta = item.get("situacao_meta")
        if situacao_meta == "cumprimento":
            row[4].text = "✓ Cumpre"
            cor_meta = VERDE
        elif situacao_meta == "incumprimento":
            row[4].text = "✗ Não cumpre"
            cor_meta = VERMELHO
        else:
            row[4].text = "— (sem meta)"
            cor_meta = CINZA
        row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shading_celula(row[4], cor_meta[0])
        for run in row[4].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(cor_meta[1])

        for cel in row:
            for run in cel.paragraphs[0].runs:
                if run.font.size is None:
                    run.font.size = Pt(10.5)
            cel.paragraphs[0].paragraph_format.space_before = Pt(5)
            cel.paragraphs[0].paragraph_format.space_after = Pt(5)

    p_legenda = doc.add_paragraph(
        "Legenda: ▲ aumento · ▼ diminuição · ● estável · — sem comparação disponível. Verde = sinal positivo "
        "· Vermelho = sinal negativo · Cinza = neutro/informativo (sem direção boa/má definida, ou sem meta "
        "aplicável). Tendência e Meta são avaliadas de forma independente uma da outra."
    )
    p_legenda.paragraph_format.space_before = Pt(6)
    p_legenda.runs[0].font.size = Pt(9)
    p_legenda.runs[0].font.color.rgb = RGBColor(0x76, 0x76, 0x76)


def indice_visual():
    doc.add_page_break()
    doc.add_heading("Índice", level=1)
    doc.add_paragraph("").paragraph_format.space_after = Pt(4)

    secoes = [
        ("1", "Governança", "1F4E79"),
        ("2", "Ambiente", "2E7D32"),
        ("3", "Mobilidade — Parque Automóvel", "6A4C93"),
        ("4", "Modos de Vida", "C0392B"),
        ("5", "Economia", "B8860B"),
        ("6", "Sociedade", "8B5E3C"),
    ]

    table = doc.add_table(rows=len(secoes), cols=2)
    table.autofit = False

    widths = (Cm(1.4), Cm(13))
    for col, w in zip(table.columns, widths):
        col.width = w

    for i, (num, nome, cor) in enumerate(secoes):
        row = table.rows[i]
        row.height = Cm(1.0)

        cel_num, cel_nome = row.cells
        for cel, w in zip((cel_num, cel_nome), widths):
            cel.width = w
            cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cel_num.text = num
        p_num = cel_num.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_num.paragraph_format.space_before = Pt(0)
        p_num.paragraph_format.space_after = Pt(0)
        run_num = p_num.runs[0]
        run_num.font.bold = True
        run_num.font.size = Pt(20)
        run_num.font.color.rgb = RGBColor.from_string(cor)

        cel_nome.text = nome
        p_nome = cel_nome.paragraphs[0]
        p_nome.paragraph_format.space_before = Pt(0)
        p_nome.paragraph_format.space_after = Pt(0)
        run_nome = p_nome.runs[0]
        run_nome.font.size = Pt(13)
        run_nome.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        _bordas_celula(cel_num, cor="D9D9D9", tamanho=4)
        _bordas_celula(cel_nome, cor="D9D9D9", tamanho=4)


# ═══════════ CAPA ═══════════
doc.add_heading("DASHBOARD CIM LEZÍRIA DO TEJO", level=1)
doc.add_heading("Relatório Automatizado", level=2)
doc.add_paragraph("Data: " + datetime.now().strftime("%d-%m-%Y"))

p_intro = doc.add_paragraph(narrativas.get("intro", ""))
p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading(f'Os 11 municípios da CIM ({tabela_dados["ano"]})', level=3)
tabela_municipios()
tabela_resumo()
indice_visual()

# ═══════════ 1. GOVERNANÇA ═══════════
titulo_seccao("1. GOVERNANÇA")

eleicoes = [
    ("aut", "Autárquicas", "01"),
    ("ar", "Legislativas (AR)", "02"),
    ("pres", "Presidenciais", "03"),
]
for codigo, nome, n in eleicoes:
    sub_titulo(f"Eleições {nome}")
    bloco_simples(
        "Taxa de Abstenção — Mapa",
        f"gov_{n}_mapa_abstenção_{codigo}",
        f"gov_{n}_mapa_abstenção_{codigo}",
        {"semEspacoAntes": True, "width": 360},
    )
    bloco_simples(
        "Taxa de Abstenção — Comparação por Município",
        f"gov_{n}_evolucao_abstenção_{codigo}",
        f"gov_{n}_evolucao_abstenção_{codigo}",
    )

sub_titulo("Telecomunicações")
bloco_simples("Acessos a Serviços de Telecomunicações", "gov_13_telecom_evolucao", "gov_13", {"semEspacoAntes": True})
bloco_simples("Indicadores de Acessibilidade Digital", "gov_14_telecom_kpis", "gov_14", {"width": 300})

# ═══════════ 2. AMBIENTE ═══════════
titulo_seccao("2. AMBIENTE")
sub_titulo("Energia")
bloco_simples("Variação Anual do Consumo", "amb_01_mapa_variacao_consumo", "amb_01", {"semEspacoAntes": True, "width": 360})
bloco_simples("Total de Contadores no Município", "amb_02_kpi_contadores", "amb_02", {"width": 220})
bloco_simples("Consumo Total de Eletricidade Anual", "amb_03_consumo_energia", "amb_03")
bloco_simples("Contadores Inteligentes", "amb_04_contadores_inteligentes", "amb_04")
bloco_simples("Consumo de Eletricidade por 1000 Habitantes", "amb_05_consumo_bt_at", "amb_05")
bloco_simples("Comunidades de Energia", "amb_06_comunidades_energia", "amb_06")
sub_titulo("Resíduos")
bloco_simples("Taxa de Deposição em Aterro", "amb_07_mapa_aterro", "amb_07", {"semEspacoAntes": True, "width": 360})
bloco_simples("Destino dos Resíduos", "amb_08_destino_residuos", "amb_08", {"width": 520})
bloco_simples("Taxa de Valorização por Município", "amb_09_valorizacao_municipio", "amb_09")

# ═══════════ 3. MOBILIDADE — PARQUE AUTOMÓVEL ═══════════
titulo_seccao("3. MOBILIDADE")
bloco_simples("Registo de Veículos Novos", "mob_01_mapa_veiculos", "mob_01", {"semEspacoAntes": True, "width": 360})
bloco_simples("Registos por Tipo de Veículo", "mob_02_registos_por_tipo", "mob_02")
bloco_simples("Evolução do Registo de Veículos", "mob_03_evolucao_veiculos", "mob_03")
bloco_simples("Pontos de Carregamento Elétrico", "mob_04_pontos_carregamento", "mob_04", {"width": 520})
bloco_simples("Carregamento Privado vs Público", "mob_05_privado_publico", "mob_05", {"width": 520})

# ═══════════ 4. MODOS DE VIDA ═══════════
titulo_seccao("4. MODOS DE VIDA")
sub_titulo("Saúde")
bloco_simples("Habitantes por Médico", "mdv_01_mapa_hab_medico", "mdv_01", {"semEspacoAntes": True, "width": 360})
bloco_simples("Profissionais de Saúde no Município", "mdv_02_kpis_profissionais", "mdv_02", {"width": 600})
bloco_simples("Utentes Inscritos no CSP", "mdv_03_kpi_utentes", "mdv_03", {"width": 220})
bloco_simples("Habitantes por Farmacêuticos e Médicos", "mdv_04_hab_farmaceuticos_medicos", "mdv_04", {"width": 600})
bloco_simples("Consultas nos Cuidados de Saúde Primários", "mdv_05_consultas_csp", "mdv_05")

sub_titulo("Segurança")
bloco_simples("Acidentes de Viação com Vítimas", "mdv_06_acidentes_viacao", "mdv_06", {"semEspacoAntes": True})
bloco_simples("Taxa de Criminalidade Total — Evolução", "mdv_07_criminalidade_evolucao", "mdv_07")
bloco_simples("Criminalidade por Tipo", "mdv_08_criminalidade_tipo", "mdv_08", {"width": 500})

sub_titulo("Educação")
bloco_simples("População Sem Nível de Escolaridade", "mdv_09_mapa_sem_escolaridade", "mdv_09", {"semEspacoAntes": True, "width": 360})
bloco_simples("Alunos Matriculados por Nível de Ensino", "mdv_10_kpis_niveis_ensino", "mdv_10", {"width": 460})
bloco_simples("Taxa de Transição/Retenção por Género", "mdv_11_transicao_retencao", "mdv_11", {"width": 480})

sub_titulo("Turismo")
bloco_simples("Indicadores de Alojamento Turístico", "mdv_12_kpis_turismo", "mdv_12", {"semEspacoAntes": True, "width": 300})
bloco_simples("Dormidas Turísticas", "mdv_13_mapa_dormidas", "mdv_13", {"width": 360})
bloco_simples("Alojamentos Vagos por Município", "mdv_14_alojamentos_vagos", "mdv_14")

# ═══════════ 5. ECONOMIA ═══════════
titulo_seccao("5. ECONOMIA")
sub_titulo("Emprego")
bloco_simples("Indicadores de Emprego (CIM)", "eco_01_kpis_emprego", "eco_01", {"semEspacoAntes": True, "width": 300})
bloco_simples("Taxa de Emprego", "eco_02_mapa_taxa_emprego", "eco_02", {"width": 360})
bloco_simples("Estrutura Setorial do Emprego", "eco_03_estrutura_setorial_emprego", "eco_03", {"width": 520})
bloco_simples("Dinâmica Empresarial", "eco_04_dinamica_empresarial", "eco_04", {"width": 560})

sub_titulo("Rendimento")
bloco_simples("Rendimento Bruto per Capita", "eco_05_mapa_rendimento", "eco_05", {"semEspacoAntes": True, "width": 360})
bloco_simples("Rendimento vs IRS per Capita", "eco_06_scatter_rendimento_irs", "eco_06", {"width": 500})
bloco_simples("Volume de Negócios per Capita", "eco_07_volume_negocios", "eco_07")
bloco_simples("Estrutura do Volume de Negócios por Setor", "eco_08_estrutura_vn", "eco_08", {"width": 520})
bloco_simples("Poder de Compra", "eco_09_poder_compra", "eco_09", {"width": 440})

# ═══════════ 6. SOCIEDADE ═══════════
titulo_seccao("6. SOCIEDADE")
bloco_simples("População Residente", "soc_01_mapa_populacao", "soc_01", {"semEspacoAntes": True, "width": 360})
bloco_simples("Evolução da População Total", "soc_02_evolucao_populacao", "soc_02")
bloco_simples("Densidade Populacional", "soc_03_mapa_densidade", "soc_03", {"width": 360})
bloco_simples("Variação Populacional por Município", "soc_04_variacao_populacional", "soc_04")
bloco_simples("População Estrangeira", "soc_05_mapa_populacao_estrangeira", "soc_05", {"width": 360})
bloco_simples("Evolução da População Estrangeira", "soc_06_evolucao_pop_estrangeira", "soc_06")
bloco_simples("População Estrangeira por Município", "soc_06b_ranking_pop_estrangeira", "soc_06b")
bloco_simples("Saldo Natural", "soc_07_saldo_natural", "soc_07")
bloco_simples("Saldo Natural por Município", "soc_07b_ranking_saldo_natural", "soc_07b")
bloco_simples("Natalidade e Mortalidade", "soc_08_natalidade_mortalidade", "soc_08")
bloco_simples("Natalidade e Mortalidade por Município", "soc_08b_ranking_natalidade_mortalidade", "soc_08b", {"width": 560})

# ── Guarda o .docx ──
out_docx = OUT_DIR / "Dashboard_CIM_Leziria.docx"
doc.save(out_docx)
print("Documento gerado com sucesso!")

# ── Conversão para PDF via LibreOffice headless ──
SOFFICE_CANDIDATOS = [
    "soffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
]

soffice_path = next((p for p in SOFFICE_CANDIDATOS if shutil.which(p) or Path(p).exists()), None)

try:
    if soffice_path is None:
        raise FileNotFoundError("soffice.exe não encontrado em nenhum caminho conhecido.")
    subprocess.run(
        [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", str(OUT_DIR), str(out_docx)],
        check=True,
        capture_output=True,
    )
    print("PDF gerado com sucesso a partir do .docx!")
except (subprocess.CalledProcessError, FileNotFoundError) as err:
    print(
        "Aviso: não foi possível gerar o PDF automaticamente (LibreOffice/soffice não "
        "encontrado ou falhou neste computador). O .docx foi gerado normalmente — para obter "
        "o PDF, abre o .docx no Word e usa 'Guardar como PDF', ou instala o LibreOffice e corre "
        f"este script novamente. Detalhe: {err}"
    )
